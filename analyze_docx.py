
import docx
from docx.shared import RGBColor
import os
import sys

def analyze_docx(path):
    doc = docx.Document(path)
    print(f"--- Analyzing {os.path.basename(path)} ---")
    
    print(f"Total Tables: {len(doc.tables)}")
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    
    print("\n--- Paragraph Dump (First 300) ---")
    for i, p in enumerate(doc.paragraphs[:300]):
        text = p.text.strip()
        if not text: continue
        
        # Check for color runs in this paragraph
        colored_text = []
        for run in p.runs:
             if run.font.color and run.font.color.rgb:
                  # detected color (usually non-black)
                  colored_text.append(f"'{run.text}'({run.font.color.rgb})")
                  
        color_info = f" [COLOR: {', '.join(colored_text)}]" if colored_text else ""
        print(f"P{i}: {text}{color_info}")

if __name__ == "__main__":
    analyze_docx("word_files/１章 時制・完了形.docx")
